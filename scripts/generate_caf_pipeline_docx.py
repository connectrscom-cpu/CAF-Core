from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def main() -> None:
    output_path = "docs/CAF_Pipeline_Automation_Structure_Depth.docx"

    doc = Document()

    title = doc.add_paragraph()
    title_run = title.add_run("CAF Pipeline — Automation, Structure, Depth")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Pipeline: Inputs Import → Evidence → Insights → Ideas → Signal Pack → Selected Ideas → "
        "Draft Package (copy/script or HeyGen package) → Rendered Assets → Human Validation → "
        "Publishing → Learning (Editorial + LLM Review + Market)"
    )

    stages: list[tuple[str, list[tuple[str, str]]]] = [
        (
            "1) Inputs Import → Evidence",
            [
                (
                    "Automation (broadness)",
                    "Continuously ingest from broad sources at scale and score everything automatically so "
                    "selection isn’t limited by manual collection or triage.",
                ),
                (
                    "Structure (observable)",
                    "Use a repeatable intake process that standardizes inputs, dedupes them, extracts features, "
                    "applies scoring, and makes a clear cutoff decision with traceability.",
                ),
                (
                    "Depth (meaning)",
                    "Evidence is useful on its own because it provides auditable ground truth and a defensible "
                    "ranking rationale for what deserves deeper analysis.",
                ),
            ],
        ),
        (
            "2) Evidence → Insights",
            [
                (
                    "Automation (broadness)",
                    "Automatically analyze every cutoff-passing evidence item, using depth tiers so deeper analysis "
                    "is reserved for the most valuable rows.",
                ),
                (
                    "Structure (observable)",
                    "Produce a consistent analysis output each time (same rubric/fields) so insights can be filtered, "
                    "grouped, compared, and reliably used downstream.",
                ),
                (
                    "Depth (meaning)",
                    "Insights are not summaries; they capture mechanisms and reusable patterns (why it worked, emotional "
                    "drivers, hook/CTA/style) that directly guide creation.",
                ),
            ],
        ),
        (
            "3) Insights → Ideas → Signal Pack",
            [
                (
                    "Automation (broadness)",
                    "Generate many ideas across audiences, formats, and platforms quickly, and produce multiple packs by "
                    "theme/campaign/persona without duplicating work.",
                ),
                (
                    "Structure (observable)",
                    "Use a systematic process to derive ideas from sets of insights and package them into a curated bundle "
                    "with ordering and intent (not a raw list).",
                ),
                (
                    "Depth (meaning)",
                    "Ideas must be standalone valuable directions (thesis, audience, why now, novelty, key points, CTA, expected "
                    "outcome) and be traceable to the insights that justify them. Signal packs must be executable and curated, "
                    "not brainstorm dumps.",
                ),
            ],
        ),
        (
            "4) Signal Pack → Selected Ideas (commitment to produce)",
            [
                (
                    "Automation (broadness)",
                    "Select the best ideas from a large pool quickly (fit to platform, campaign, audience; diversity; urgency) "
                    "without needing deep manual review of everything.",
                ),
                (
                    "Structure (observable)",
                    "Create a clear selection outcome: an ordered list of ideas committed for production, with lightweight selection "
                    "reasons and preserved traceability.",
                ),
                (
                    "Depth (meaning)",
                    "Selection output is a production plan, not an idea list—each selected idea is specific enough that downstream "
                    "work is execution, not interpretation.",
                ),
            ],
        ),
        (
            "5A) Selected Ideas → Draft Package (execution-ready copy OR execution-ready HeyGen package)",
            [
                (
                    "Automation (broadness)",
                    "Generate many execution-ready packages in parallel and create meaningful variants when they improve outcomes "
                    "(different hook strategy, structure, tone), not random paraphrases.",
                ),
                (
                    "Structure (observable)",
                    "Produce strict, tool-ready outputs: either render-ready copy (including slide-by-slide / on-screen text where "
                    "relevant) or a video-generation package (script + scene plan + delivery guidance).",
                ),
                (
                    "Depth (meaning)",
                    "Draft packages must be ready to run with minimal edits—complete, specific, aligned to the selected idea’s thesis/key "
                    "points, and grounded (not vague drafts).",
                ),
            ],
        ),
        (
            "5B) Draft Package → Rendered Assets (template render OR HeyGen execution)",
            [
                (
                    "Automation (broadness)",
                    "Execute rendering at scale (multiple variants/aspect ratios/templates) with reliable throughput and safe retries.",
                ),
                (
                    "Structure (observable)",
                    "Use a repeatable production step that consistently turns packages into concrete deliverables with clear tracking of what "
                    "was produced.",
                ),
                (
                    "Depth (meaning)",
                    "Outputs are real, reviewable media—presentation-grade, readable, brand-consistent, coherent pacing/captions—ready for "
                    "human review.",
                ),
            ],
        ),
        (
            "6) Human Validation (reviewed content + improvement guidance + editorial labels)",
            [
                (
                    "Automation (broadness)",
                    "Reduce bottlenecks via routing and structured feedback capture so learning scales across all reviewed items (approved or not).",
                ),
                (
                    "Structure (observable)",
                    "Use a consistent review workflow that records verdict, what’s wrong and where, specific suggestions/example fixes, and standardized labels.",
                ),
                (
                    "Depth (meaning)",
                    "Validation output is not just approval; it’s reviewed content plus actionable improvement guidance and diagnosis. This becomes the highest-quality "
                    "input for editorial learning.",
                ),
            ],
        ),
        (
            "7) Publishing (validated → distributed with attribution)",
            [
                (
                    "Automation (broadness)",
                    "Publish/schedule across platforms efficiently, handle variants, and log what shipped (where/when) and what’s needed for tracking without manual reporting.",
                ),
                (
                    "Structure (observable)",
                    "Use a consistent release process that records what shipped, where, when, and the intended goal/hypothesis so outcomes can be evaluated.",
                ),
                (
                    "Depth (meaning)",
                    "Publishing is treated as an experiment: each piece has intent and is attributable back to the upstream idea/insights that generated it.",
                ),
            ],
        ),
        (
            "8) Learning (three loops)",
            [
                (
                    "8A) Editorial Analysis — Automation",
                    "Aggregate patterns across all human-reviewed content to identify what consistently passes/fails and where effort is wasted.",
                ),
                (
                    "8A) Editorial Analysis — Structure",
                    "Systematically convert review feedback into prioritized findings, operational fixes, and generation guidance.",
                ),
                (
                    "8A) Editorial Analysis — Depth",
                    "Produce implementable improvements: what to boost/avoid by approval rate, which templates need fixes, and what rules/guidance should change upstream.",
                ),
                (
                    "8B) LLM Review — Automation",
                    "Review approved outputs at scale and generate concrete improvement suggestions and alternatives.",
                ),
                (
                    "8B) LLM Review — Structure",
                    "Standardize suggestions so they’re comparable and reusable, and can be adopted as guidance if desired.",
                ),
                (
                    "8B) LLM Review — Depth",
                    "Suggestions must be specific and applicable (rewrites/structural upgrades), not vague commentary.",
                ),
                (
                    "8C) Market Analysis — Automation",
                    "Continuously collect performance signals and translate them into recommendations without manual interpretation.",
                ),
                (
                    "8C) Market Analysis — Structure",
                    "Normalize results and attribute performance back to controllable levers (formats, themes, hooks, CTAs, templates, insight patterns).",
                ),
                (
                    "8C) Market Analysis — Depth",
                    "Produce clear boost/derank and opportunity allocation guidance, grounded in outcomes and traceable back to upstream decisions.",
                ),
            ],
        ),
    ]

    for stage_title, bullets in stages:
        p = doc.add_paragraph(stage_title)
        p.runs[0].bold = True
        for head, text in bullets:
            para = doc.add_paragraph()
            r = para.add_run(f"{head}: ")
            r.bold = True
            para.add_run(text)

    p = doc.add_paragraph()
    r = p.add_run("Guiding principle")
    r.bold = True
    p.add_run(
        ": Every stage produces an output that is useful on its own—automation expands scale, structure makes the "
        "process observable and repeatable, and depth ensures outputs are meaningful (not fluff)."
    )

    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()

